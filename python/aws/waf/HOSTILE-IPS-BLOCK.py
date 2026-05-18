import boto3
import requests
import sys
from docx import Document

# ─────────────────────────────────────────────
# Prompt for all sensitive values at runtime
# ─────────────────────────────────────────────
print("=" * 55)
print("  AWS WAF - Hostile IP Blocker")
print("=" * 55)

API_KEY     = input("Enter your AbuseIPDB API key: ").strip()
WAF_ARN     = input("Enter your WAF WebACL ARN: ").strip()
IPSET_ARN   = input("Enter your WAF IP Set ARN: ").strip()
WAF_REGION  = input("Enter AWS region (e.g. eu-west-1): ").strip()

if not all([API_KEY, WAF_ARN, IPSET_ARN, WAF_REGION]):
    print("ERROR: All fields are required.")
    sys.exit(1)

# ─────────────────────────────────────────────
# AWS WAF client
# ─────────────────────────────────────────────
waf_client = boto3.client('wafv2', region_name=WAF_REGION)

# ─────────────────────────────────────────────
# Get current IPs in the IP Set
# ─────────────────────────────────────────────
def get_ips_in_ipset(ipset_arn):
    name = ipset_arn.split('/')[-2]
    id_  = ipset_arn.split('/')[-1]
    response = waf_client.get_ip_set(Name=name, Scope='REGIONAL', Id=id_)
    return response['IPSet']['Addresses'], response['LockToken']

# ─────────────────────────────────────────────
# Add hostile IPs to WAF IP Set (WAFv2 correct API)
# ─────────────────────────────────────────────
def add_ips_to_ipset(ipset_arn, ips_to_add):
    if not ips_to_add:
        print("No new IPs to add.")
        return

    name      = ipset_arn.split('/')[-2]
    id_       = ipset_arn.split('/')[-1]
    current_ips, lock_token = get_ips_in_ipset(ipset_arn)

    new_ips = list(set(current_ips + ips_to_add))

    waf_client.update_ip_set(
        Name=name,
        Scope='REGIONAL',
        Id=id_,
        Addresses=new_ips,
        LockToken=lock_token
    )
    print(f"Added {len(ips_to_add)} IPs to WAF IP Set.")

# ─────────────────────────────────────────────
# Check IP against AbuseIPDB
# ─────────────────────────────────────────────
def check_ip(ip):
    url = 'https://api.abuseipdb.com/api/v2/check'
    headers = {'Key': API_KEY, 'Accept': 'application/json'}
    params  = {'ipAddress': ip, 'maxAgeInDays': '90'}
    try:
        response = requests.get(url, headers=headers, params=params)
        if response.status_code == 200:
            score = response.json().get('data', {}).get('abuseConfidenceScore')
            print(f"  {ip} → score: {score}")
            return score
        else:
            print(f"  {ip} → HTTP {response.status_code}")
    except Exception as e:
        print(f"  {ip} → Error: {e}")
    return None

# ─────────────────────────────────────────────
# Prompt for IPs to check
# ─────────────────────────────────────────────
print("\nHow would you like to provide IPs to check?")
print("  1. Load from text file (one IP per line)")
print("  2. Enter manually (comma-separated)")
choice = input("Choice (1/2): ").strip()

ips_to_check = []
if choice == "1":
    ip_file = input("Path to IP file: ").strip()
    try:
        with open(ip_file) as f:
            ips_to_check = [line.strip() for line in f if line.strip()]
    except FileNotFoundError:
        print(f"ERROR: File '{ip_file}' not found.")
        sys.exit(1)
elif choice == "2":
    raw = input("Enter IPs (comma-separated): ").strip()
    ips_to_check = [ip.strip() for ip in raw.split(",") if ip.strip()]
else:
    print("Invalid choice.")
    sys.exit(1)

# ─────────────────────────────────────────────
# Analyze IPs
# ─────────────────────────────────────────────
print(f"\nChecking {len(ips_to_check)} IPs against AbuseIPDB...")
hostile_ips    = []
legitimate_ips = []
ips_to_block   = []

current_ips_in_set, _ = get_ips_in_ipset(IPSET_ARN)

for ip in ips_to_check:
    score = check_ip(ip)
    if score is not None:
        cidr = f"{ip}/32"
        if score >= 80:
            hostile_ips.append(cidr)
            if cidr not in current_ips_in_set:
                ips_to_block.append(cidr)
        else:
            legitimate_ips.append(cidr)

# ─────────────────────────────────────────────
# Add to WAF
# ─────────────────────────────────────────────
add_ips_to_ipset(IPSET_ARN, ips_to_block)

# ─────────────────────────────────────────────
# Word report
# ─────────────────────────────────────────────
doc = Document()
doc.add_heading('IP Address Analysis from AWS WAF', 0)

doc.add_heading('Hostile IPs (score >= 80)', level=1)
for ip in hostile_ips:
    doc.add_paragraph(ip)

doc.add_heading('Legitimate IPs (score < 80)', level=1)
for ip in legitimate_ips:
    doc.add_paragraph(ip)

output_file = 'IP_Address_Analysis_with_subnet_from_WAF.docx'
doc.save(output_file)
print(f"\nReport saved to: {output_file}")
print(f"Hostile: {len(hostile_ips)} | Legitimate: {len(legitimate_ips)} | Added to WAF: {len(ips_to_block)}")
