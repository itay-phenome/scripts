import requests
import sys
from docx import Document

# ─────────────────────────────────────────────
# Prompt for API key at runtime (never hardcode!)
# ─────────────────────────────────────────────
print("=" * 50)
print("  AbuseIPDB - IP Check Script")
print("=" * 50)
API_KEY = input("Enter your AbuseIPDB API key: ").strip()
if not API_KEY:
    print("ERROR: API key cannot be empty.")
    sys.exit(1)

# ─────────────────────────────────────────────
# Prompt for IP list — file or manual input
# ─────────────────────────────────────────────
print("\nHow would you like to provide the IP list?")
print("  1. Load from a text file (one IP per line)")
print("  2. Enter IPs manually (comma-separated)")
choice = input("Choice (1/2): ").strip()

ips_to_check = []

if choice == "1":
    ip_file = input("Enter path to IP file: ").strip()
    try:
        with open(ip_file, "r") as f:
            ips_to_check = [line.strip() for line in f if line.strip()]
        print(f"Loaded {len(ips_to_check)} IPs from {ip_file}")
    except FileNotFoundError:
        print(f"ERROR: File '{ip_file}' not found.")
        sys.exit(1)
elif choice == "2":
    raw = input("Enter IPs separated by commas: ").strip()
    ips_to_check = [ip.strip() for ip in raw.split(",") if ip.strip()]
    print(f"Loaded {len(ips_to_check)} IPs.")
else:
    print("Invalid choice. Exiting.")
    sys.exit(1)

if not ips_to_check:
    print("ERROR: No IPs to check.")
    sys.exit(1)

# ─────────────────────────────────────────────
# Check each IP against AbuseIPDB
# ─────────────────────────────────────────────
def check_ip(ip):
    url = 'https://api.abuseipdb.com/api/v2/check'
    headers = {
        'Key': API_KEY,
        'Accept': 'application/json',
    }
    params = {
        'ipAddress': ip,
        'maxAgeInDays': '90'
    }
    try:
        response = requests.get(url, headers=headers, params=params)
        print(f"Checking: {ip}")
        if response.status_code == 200:
            data = response.json()
            abuse_score = data.get('data', {}).get('abuseConfidenceScore', None)
            if abuse_score is not None:
                return abuse_score
            else:
                print(f"  No abuse score found for {ip}")
        else:
            print(f"  Failed for {ip} — HTTP {response.status_code}")
    except Exception as e:
        print(f"  Error checking {ip}: {e}")
    return None

hostile_ips = []
legitimate_ips = []

for ip in ips_to_check:
    score = check_ip(ip)
    if score is not None:
        if score >= 80:
            hostile_ips.append(ip)
        else:
            legitimate_ips.append(ip)

# ─────────────────────────────────────────────
# Generate Word report
# ─────────────────────────────────────────────
doc = Document()
doc.add_heading('IP Address Analysis', 0)

doc.add_heading('Hostile IPs (score >= 80)', level=1)
if hostile_ips:
    for ip in hostile_ips:
        doc.add_paragraph(ip)
else:
    doc.add_paragraph("None found.")

doc.add_heading('Legitimate IPs (score < 80)', level=1)
if legitimate_ips:
    for ip in legitimate_ips:
        doc.add_paragraph(ip)
else:
    doc.add_paragraph("None found.")

output_file = 'IP_Address_Analysis.docx'
doc.save(output_file)

print(f"\nDone! Report saved to: {output_file}")
print(f"Hostile: {len(hostile_ips)} | Legitimate: {len(legitimate_ips)}")
